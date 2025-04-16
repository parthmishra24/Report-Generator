from docx import Document
from docx.shared import Inches
import os
import sys
import datetime
import questionary
from style import custom_style  # ‚úÖ Import your centralized style

def generate_docx_report(vulnerabilities, template_path, output_path):
    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"‚ùå Error loading template: {e}")
        return

    # Auto-date in DD/MM/YYYY
    report_date = datetime.datetime.today().strftime("%d/%m/%Y")

    # Ask for target URL
    target_url = questionary.text("üåê Enter the target URL:", style=custom_style).ask()

    # Replace placeholders
    for paragraph in doc.paragraphs:
        if "{Date}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{Date}", report_date)
        if "{Target URL}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{Target URL}", target_url)

    # Section 5 ‚Äì Vulnerability Table
    doc.add_page_break()
    doc.add_heading("5. Vulnerability finding table", level=1)

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'S.No'
    hdr_cells[1].text = 'CWE'
    hdr_cells[2].text = 'Vulnerability'
    hdr_cells[3].text = 'Severity'
    hdr_cells[4].text = 'Status'

    for i, vuln in enumerate(vulnerabilities, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = vuln['cwe_id']
        row_cells[2].text = vuln['name']
        row_cells[3].text = vuln['severity']
        row_cells[4].text = vuln['status']

    # Section 6 ‚Äì Vulnerability Details
    doc.add_page_break()
    doc.add_heading("6. Vulnerability Finding Details", level=1)

    for idx, vuln in enumerate(vulnerabilities, start=1):
        doc.add_heading(f"6.{idx} {vuln['name']}", level=2)

        doc.add_paragraph("Description", style='Heading 3')
        doc.add_paragraph(vuln['description'])

        doc.add_paragraph("Affected URL", style='Heading 3')
        doc.add_paragraph(vuln['affected_url'])

        doc.add_paragraph("Impact", style='Heading 3')
        doc.add_paragraph(vuln['impact'])

        doc.add_paragraph("Remediation", style='Heading 3')
        doc.add_paragraph(vuln['remediation'])

        doc.add_paragraph("Evidence", style='Heading 3')
        if vuln.get('screenshots'):
            for path in vuln['screenshots']:
                if os.path.exists(path):
                    try:
                        doc.add_picture(path, width=Inches(5.5))
                    except Exception as e:
                        doc.add_paragraph(f"‚ö†Ô∏è Could not add screenshot '{path}': {e}")
                else:
                    doc.add_paragraph(f"‚ö†Ô∏è Screenshot not found: {path}")
        else:
            doc.add_paragraph("‚ö†Ô∏è No screenshots provided.")

        doc.add_page_break()

    # Ask for file name
    filename_input = questionary.text("üìÑ Enter custom report file name (without extension):", style=custom_style).ask()
    if filename_input and filename_input.strip():
        filename = f"{filename_input.strip()}.docx"
    else:
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"VAPT_Report_{timestamp}.docx"

    save_path = os.path.join(output_path, filename)

    try:
        doc.save(save_path)
        print(f"‚úÖ DOCX report saved to: {save_path}")
    except Exception as e:
        print(f"‚ùå Failed to save report: {e}")

    # Ask user if they want to open the report
    open_report = questionary.confirm("üìÇ Do you want to open the generated report?", style=custom_style).ask()
    if open_report:
        try:
            import subprocess
            if os.name == 'nt':
                os.startfile(save_path)
            elif os.name == 'posix':
                subprocess.run(['open' if sys.platform == 'darwin' else 'xdg-open', save_path])
        except Exception as e:
            print(f"‚ö†Ô∏è Could not open the report: {e}")
